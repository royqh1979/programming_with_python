from docx import Document
from easygraphics import dialog as dlg

import os
from pathlib import Path

package_count = 0
def gen_word(doc,path:Path, package_name:str):
    global  package_count
    package_count+=1
    doc.add_heading(str(package_count)+" "+package_name+"模组", level=1)
    file_count = 0
    blank_line_count = 0
    for entry in path.iterdir():
        if not entry.is_file():
            continue
        filename = str(path.absolute()) + os.sep + entry.name
        if filename.endswith(".py") :
            print(filename)
            file_count += 1
            doc.add_heading(str(package_count)+"."+str(file_count)+" 文件"+entry.name,level=2)
            with open(filename,"r",encoding="UTF-8") as file:
                for line in file.readlines():
                    line = line.rstrip()
                    if len(line)==0:
                        blank_line_count+=1
                    else:
                        blank_line_count = 0
                    if blank_line_count>1:
                        continue
                    doc.add_paragraph(line)
    for entry in path.iterdir():
        if not entry.is_dir():
            continue
        if entry.name.startswith("__"):
            continue
        gen_word(doc,entry,package_name+"."+entry.name)



dir_path = dlg.get_directory_name("请选择要转换的目录")

path = Path(dir_path)

doc = Document()

gen_word(doc,path,path.name)

save_filename = dlg.get_save_file_name("请选择要保存到的文件")
doc.save(save_filename)